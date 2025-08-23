#!/usr/bin/env python3
"""
Document Conversion Script for ERCOT Nodal Protocols
Converts .docx and .doc files to searchable text format
"""

import os
import sys
import subprocess
from pathlib import Path
import time
from datetime import datetime

# Track progress
class DocumentConverter:
    def __init__(self, base_dir="."):
        self.base_dir = Path(base_dir)
        self.text_dir = self.base_dir / "text_versions"
        self.text_dir.mkdir(exist_ok=True)
        
        self.total_files = 0
        self.converted_files = 0
        self.failed_files = []
        self.skipped_files = []
        
    def get_documents(self):
        """Get list of all documents to convert"""
        docx_files = list(self.base_dir.glob("*.docx"))
        doc_files = list(self.base_dir.glob("*.doc"))
        all_files = docx_files + doc_files
        
        # Filter out already converted files
        files_to_convert = []
        for file in all_files:
            txt_name = file.stem + ".txt"
            txt_path = self.base_dir / txt_name
            if not txt_path.exists():
                files_to_convert.append(file)
            else:
                self.skipped_files.append(file.name)
                
        return sorted(files_to_convert)
    
    def convert_with_libreoffice(self, input_file):
        """Convert document using LibreOffice"""
        try:
            # Use LibreOffice in headless mode to convert to text
            cmd = [
                "libreoffice",
                "--headless",
                "--convert-to",
                "txt:Text",
                "--outdir",
                str(self.base_dir),
                str(input_file)
            ]
            
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=60
            )
            
            if result.returncode == 0:
                return True
            else:
                print(f"  Error: {result.stderr}")
                return False
                
        except subprocess.TimeoutExpired:
            print(f"  Timeout converting {input_file.name}")
            return False
        except Exception as e:
            print(f"  Error converting {input_file.name}: {e}")
            return False
    
    def convert_with_python_docx(self, input_file):
        """Convert .docx using python-docx library"""
        try:
            # First try to import python-docx
            import docx
            
            doc = docx.Document(str(input_file))
            
            # Extract all text from paragraphs and tables
            full_text = []
            
            # Get paragraph text
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Get table text
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append('\t'.join(row_text))
            
            # Write to text file
            output_file = self.base_dir / (input_file.stem + ".txt")
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write('\n'.join(full_text))
            
            return True
            
        except ImportError:
            # python-docx not installed, fall back to LibreOffice
            return self.convert_with_libreoffice(input_file)
        except Exception as e:
            print(f"  Python-docx failed: {e}, trying LibreOffice...")
            return self.convert_with_libreoffice(input_file)
    
    def convert_document(self, file_path):
        """Convert a single document to text"""
        print(f"Converting: {file_path.name}")
        
        # Try python-docx for .docx files, LibreOffice for all
        if file_path.suffix.lower() == '.docx':
            success = self.convert_with_python_docx(file_path)
        else:
            success = self.convert_with_libreoffice(file_path)
        
        if success:
            self.converted_files += 1
            print(f"  ✓ Converted successfully")
            
            # Move to text_versions directory if desired
            # txt_file = self.base_dir / (file_path.stem + ".txt")
            # if txt_file.exists():
            #     txt_file.rename(self.text_dir / txt_file.name)
        else:
            self.failed_files.append(file_path.name)
            print(f"  ✗ Conversion failed")
        
        return success
    
    def run_conversion(self):
        """Run the batch conversion process"""
        print("=" * 60)
        print("ERCOT Document Conversion Tool")
        print("=" * 60)
        print(f"Working directory: {self.base_dir.absolute()}")
        print(f"Output directory: {self.text_dir.absolute()}")
        print()
        
        # Get documents to convert
        documents = self.get_documents()
        self.total_files = len(documents)
        
        if self.skipped_files:
            print(f"Skipping {len(self.skipped_files)} already converted files:")
            for f in self.skipped_files[:5]:
                print(f"  - {f}")
            if len(self.skipped_files) > 5:
                print(f"  ... and {len(self.skipped_files) - 5} more")
            print()
        
        if not documents:
            print("No documents to convert!")
            return
        
        print(f"Found {self.total_files} documents to convert")
        print("-" * 60)
        
        start_time = time.time()
        
        # Convert each document
        for i, doc in enumerate(documents, 1):
            print(f"\n[{i}/{self.total_files}] ", end="")
            self.convert_document(doc)
        
        # Print summary
        elapsed_time = time.time() - start_time
        print("\n" + "=" * 60)
        print("CONVERSION SUMMARY")
        print("=" * 60)
        print(f"Total documents: {self.total_files}")
        print(f"Successfully converted: {self.converted_files}")
        print(f"Failed conversions: {len(self.failed_files)}")
        print(f"Already converted: {len(self.skipped_files)}")
        print(f"Time elapsed: {elapsed_time:.1f} seconds")
        
        if self.failed_files:
            print("\nFailed files:")
            for f in self.failed_files:
                print(f"  - {f}")
        
        print("\n✓ Conversion process complete!")
        
        # Create conversion log
        self.create_log()
    
    def create_log(self):
        """Create a log file of the conversion process"""
        log_file = self.base_dir / "conversion_log.txt"
        
        with open(log_file, 'w') as f:
            f.write(f"ERCOT Document Conversion Log\n")
            f.write(f"Generated: {datetime.now().isoformat()}\n")
            f.write(f"{'=' * 60}\n\n")
            
            f.write(f"Summary:\n")
            f.write(f"  Total files: {self.total_files}\n")
            f.write(f"  Converted: {self.converted_files}\n")
            f.write(f"  Failed: {len(self.failed_files)}\n")
            f.write(f"  Skipped: {len(self.skipped_files)}\n\n")
            
            if self.failed_files:
                f.write(f"Failed conversions:\n")
                for file in self.failed_files:
                    f.write(f"  - {file}\n")
                f.write("\n")
            
            if self.skipped_files:
                f.write(f"Already converted (skipped):\n")
                for file in self.skipped_files:
                    f.write(f"  - {file}\n")


def main():
    # Check if we're in the right directory
    current_dir = Path.cwd()
    
    # Look for marker files to confirm we're in the right place
    if not any(current_dir.glob("*Nodal.docx")) and not any(current_dir.glob("*Nodal.doc")):
        print("Error: No ERCOT Nodal Protocol documents found in current directory")
        print(f"Current directory: {current_dir}")
        sys.exit(1)
    
    # Check for LibreOffice
    try:
        result = subprocess.run(["libreoffice", "--version"], capture_output=True)
        if result.returncode != 0:
            print("Warning: LibreOffice not found or not working properly")
            print("Some conversions may fail")
    except FileNotFoundError:
        print("Warning: LibreOffice not installed")
        print("Install with: sudo apt-get install libreoffice")
        response = input("Continue anyway? (y/n): ")
        if response.lower() != 'y':
            sys.exit(1)
    
    # Try to install python-docx if not available
    try:
        import docx
    except ImportError:
        print("python-docx not installed. Attempting to install...")
        subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"], 
                      capture_output=True)
    
    # Run conversion
    converter = DocumentConverter(current_dir)
    converter.run_conversion()


if __name__ == "__main__":
    main()